"""
Generates SUBMISSION_REPORT.docx — ~20 page university submission document.
Run: python make_report.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ───────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Colours ────────────────────────────────────────────────────────────────────
DARK_BLUE = RGBColor(0x1F, 0x38, 0x64)
MID_BLUE  = RGBColor(0x2F, 0x54, 0x96)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
GREY      = RGBColor(0x44, 0x44, 0x44)

# ── Helpers ────────────────────────────────────────────────────────────────────
def page_break():
    doc.add_page_break()

def spacer(pt=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pt)
    p.paragraph_format.space_before = Pt(0)

def h1(text):
    p = doc.add_paragraph()
    p.clear()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = DARK_BLUE
    run.font.size = Pt(16)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(8)
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F3864')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def h2(text):
    p = doc.add_paragraph()
    p.clear()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = DARK_BLUE
    run.font.size = Pt(13)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    return p

def h3(text):
    p = doc.add_paragraph()
    p.clear()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = MID_BLUE
    run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    return p

def h4(text):
    p = doc.add_paragraph()
    p.clear()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = MID_BLUE
    run.font.size = Pt(10.5)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    return p

def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def numbered(text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    return p

def quote_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.right_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10.5)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EBF3FB')
    pPr.append(shd)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F3864')
        tcPr.append(shd)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            for para in cells[ci].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
            if (ri % 2) == 1:
                tcPr2 = cells[ci]._tc.get_or_add_tcPr()
                shd2 = OxmlElement('w:shd')
                shd2.set(qn('w:val'), 'clear')
                shd2.set(qn('w:color'), 'auto')
                shd2.set(qn('w:fill'), 'EBF3FB')
                tcPr2.append(shd2)
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    spacer(8)

def info_row(label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(label + ': ')
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = DARK_BLUE
    r2 = p.add_run(value)
    r2.font.size = Pt(11)

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 1 — COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
spacer(60)

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run('CBD Assignment 2')
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = DARK_BLUE

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub_p.add_run('AI-Assisted Test Development Report')
r2.font.size = Pt(18)
r2.font.color.rgb = MID_BLUE

spacer(4)

app_p = doc.add_paragraph()
app_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = app_p.add_run('Task Manager REST API  ·  Spring Boot 3.2')
r3.font.size = Pt(13)
r3.font.color.rgb = GREY

spacer(40)

for label, value in [
    ('Student Name',  'Kirthika Alagar'),
    ('Student ID',    'A00336252'),
    ('Course',        'MSc Software Design with Cloud Native Computing'),
    ('Module',        'Cloud-Based Development (CBD)'),
    ('Submission Date', '26 March 2025'),
    ('Weighting',     '20%'),
    ('GitHub',        'https://github.com/CodingKirthi/CBD-Taskmanager'),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    rb = p.add_run(label + ':  ')
    rb.bold = True
    rb.font.size = Pt(11)
    rb.font.color.rgb = DARK_BLUE
    rv = p.add_run(value)
    rv.font.size = Pt(11)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 2 — TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════════════════════
h1('Table of Contents')
spacer(4)

toc_items = [
    ('Introduction', '3'),
    ('Application Under Test', '3'),
    ('Section 1 – Test Strategy', '4'),
    ('  1.1  Alignment with the Test Pyramid', '4'),
    ('  1.2  Test Pyramid Distribution', '5'),
    ('Section 2 – Dependency Injection and Mocking Approach', '6'),
    ('  2.1  How Constructor Injection Enables Testability', '6'),
    ('  2.2  Where Mocks Were Used and Why', '7'),
    ('  2.3  Where Mocks Were Avoided and Why', '7'),
    ('Section 3 – Mapping to Test Pyramid', '8'),
    ('  3.1  Unit Tests – 15 tests', '8'),
    ('  3.2  Integration Tests – 9 tests', '10'),
    ('  3.3  End-to-End Tests – 4 tests', '11'),
    ('Section 4 – AI Interaction Log', '12'),
    ('  4.1  Example 1: AI Output Accepted As-Is', '12'),
    ('  4.2  Example 2: AI Output Modified', '13'),
    ('  4.3  Example 3: AI Output Rejected', '14'),
    ('  4.4  Example 4: AI-Identified Gap', '15'),
    ('  4.5  Example 5: AI Output Constrained by Explicit Rules', '16'),
    ('Section 5 – Evaluation and Reflection', '17'),
    ('Conclusion', '19'),
    ('References', '20'),
]
for title, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(14.5), 2)   # right-align at 14.5 cm
    indent = title.startswith('  ')
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
        r_title = p.add_run(title.strip())
        r_title.font.size = Pt(10.5)
        r_title.font.color.rgb = MID_BLUE
    else:
        r_title = p.add_run(title)
        r_title.bold = True
        r_title.font.size = Pt(11)
        r_title.font.color.rgb = DARK_BLUE
    r_tab = p.add_run('\t' + page)
    r_tab.font.size = Pt(10.5)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 3 — INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
h1('Introduction')
body(
    'This report documents the design, implementation, and evaluation of a comprehensive automated test suite '
    'for a Task Manager REST API built with Spring Boot 3.2. The report fulfils the requirements of CBD '
    'Assignment 2, which requires the use of Artificial Intelligence (AI) assistance during the test development '
    'process and a critical reflection on that experience.'
)
body(
    'The assignment covers three primary areas: (1) designing a test strategy aligned with the Test Pyramid, '
    '(2) applying appropriate dependency injection and mocking techniques to isolate components under test, '
    'and (3) documenting interactions with an AI assistant — recording where AI output was accepted, modified, '
    'rejected, or constrained.'
)
body(
    'The goal of this assignment is not simply to achieve passing tests, but to demonstrate a principled '
    'approach to software quality — one in which each test is placed at the correct level of the pyramid, '
    'tests the right thing in the right way, and can be relied upon to catch real regressions without '
    'introducing unnecessary coupling to implementation details or infrastructure.'
)
body(
    'All source code, tests, and configuration files are available in the GitHub repository at: '
    'https://github.com/CodingKirthi/CBD-Taskmanager'
)

spacer(8)

# ══════════════════════════════════════════════════════════════════════════════
# APPLICATION UNDER TEST
# ══════════════════════════════════════════════════════════════════════════════
h1('Application Under Test')
body(
    'The application is a Task Manager REST API built with Spring Boot 3.2.3. It exposes RESTful HTTP '
    'endpoints for managing two core domain objects: Projects and Tasks. The API enforces the following '
    'business rules at the service layer:'
)
bullet('A task cannot be created with a due date in the past.')
bullet('A completed task (status = DONE) cannot be reopened to status = OPEN.')
bullet('A project cannot be deleted while any of its tasks have status = IN_PROGRESS.')
bullet('Project names must be unique across the system.')

spacer(4)
h3('Technology Stack')
add_table(
    ['Component', 'Technology / Version'],
    [
        ['Web Framework',     'Spring Boot 3.2.3 (Spring MVC)'],
        ['Persistence',       'Spring Data JPA + Hibernate'],
        ['Database (Runtime)','H2 file-based database'],
        ['Database (Tests)',  'H2 in-memory database'],
        ['Validation',        'Jakarta Bean Validation 3.0'],
        ['Test Framework',    'JUnit 5 (JUnit Jupiter)'],
        ['Mocking',           'Mockito 5'],
        ['Assertions',        'AssertJ 3'],
        ['Build Tool',        'Apache Maven'],
        ['Java Version',      'Java 17'],
    ],
    col_widths=[5, 11]
)

h3('Domain Model')
body(
    'The domain consists of two entities. A Project has a name, an optional description, and a collection '
    'of Tasks. A Task belongs to exactly one Project and has a title, optional description, status '
    '(OPEN / IN_PROGRESS / DONE), priority (LOW / MEDIUM / HIGH / CRITICAL), an optional due date, '
    'and an optional completedAt timestamp set automatically when status transitions to DONE.'
)
body(
    'REST endpoints are exposed through TaskController and ProjectController. Business logic is encapsulated '
    'in TaskService and ProjectService. Data access is handled by TaskRepository and ProjectRepository, '
    'which extend JpaRepository. TaskMapper converts Task entities to TaskDTO objects for the API response.'
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — TEST STRATEGY
# ══════════════════════════════════════════════════════════════════════════════
h1('Section 1 – Test Strategy')

h2('1.1  Alignment with the Test Pyramid')
body(
    'The test suite is intentionally structured in three tiers that mirror the Test Pyramid, a concept '
    'introduced by Mike Cohn and refined by Martin Fowler. The pyramid prescribes that a healthy test suite '
    'should have many fast, cheap tests at the bottom (unit), fewer slower tests in the middle (integration), '
    'and very few comprehensive but expensive tests at the top (end-to-end). This distribution minimises the '
    'total feedback time while maximising the reliability of the suite.'
)

code_block(
    "            /\\\n"
    "           /  \\        E2E (4 tests)  – Full HTTP stack, real H2 DB\n"
    "          /    \\\n"
    "         /------\\      Integration (9 tests)  – JPA slice + Web MVC slice\n"
    "        /        \\\n"
    "       /----------\\    Unit (15 tests)  – No Spring context, Mockito only\n"
    "      /______________\\"
)

h4('Unit Tests  (15 tests, ~0.7 s total run time)')
body(
    'Unit tests form the base of the pyramid (54% of the suite). They test business logic in complete '
    'isolation using @ExtendWith(MockitoExtension.class). No Spring context is ever started. Every dependency '
    'of the class under test is replaced with a Mockito @Mock, and the class itself is instantiated via '
    '@InjectMocks. This means each test runs in milliseconds and fails only when the business logic itself '
    'is wrong — not because of database state, HTTP routing, or bean wiring.'
)
body(
    'Unit tests cover: TaskService (8 tests covering all public methods and their guard clauses), '
    'ProjectService (5 tests covering creation, deletion constraints, and retrieval), and TaskMapper '
    '(2 tests verifying the entity-to-DTO mapping contract).'
)

h4('Integration Tests  (9 tests, ~3 s total run time)')
body(
    'Integration tests occupy the middle tier (32% of the suite). They operate at two distinct Spring '
    'test slices — a narrow subset of the full application context:'
)
bullet(
    '@DataJpaTest (4 tests) — Loads only the JPA slice: repositories, the entity manager, and an H2 '
    'in-memory schema. No web layer or service layer is present. These tests verify that derived query '
    'methods (findByProjectId, findByStatus, countByProjectIdAndStatus, findByPriorityAndStatusNot) '
    'produce correct SQL against a real schema.'
)
bullet(
    '@WebMvcTest (5 tests) — Loads only the web layer: controllers, the GlobalExceptionHandler, and bean '
    'validation. All service beans are replaced with @MockBean. These tests verify HTTP method mapping, '
    'request deserialization, bean validation enforcement, and correct HTTP status code mapping.'
)

h4('End-to-End Tests  (4 tests, ~11 s total run time)')
body(
    'End-to-end tests sit at the apex of the pyramid (14% of the suite). They use '
    '@SpringBootTest(webEnvironment = RANDOM_PORT) to boot the complete application, and TestRestTemplate '
    'to make real HTTP calls over the network stack. An H2 in-memory database serves as the persistence '
    'layer, providing isolation between test runs. These tests exercise the full request → controller → '
    'service → repository → response pipeline, giving the highest confidence that the system works as a whole.'
)

spacer(4)
h2('1.2  Test Pyramid Distribution')
add_table(
    ['Level', 'Test Count', '% of Suite', 'Approx. Run Time', 'Spring Context'],
    [
        ['Unit',         '15', '54%', '< 1 ms each',    'None (Mockito only)'],
        ['Integration',  '9',  '32%', '50–500 ms each', '@DataJpaTest / @WebMvcTest'],
        ['End-to-End',   '4',  '14%', '1–3 s each',     '@SpringBootTest (full)'],
        ['Total',        '28', '100%', '~15 s total',    ''],
    ],
    col_widths=[3.5, 2.5, 2.5, 3.5, 5.5]
)

body(
    'The 54/32/14 split closely matches the recommended Test Pyramid proportion. The relatively small '
    'number of E2E tests is intentional: each one is slow (full Spring startup + HTTP round-trip), '
    'potentially flaky (relies on network stack, port allocation), and expensive to maintain. Unit tests '
    'provide fast, deterministic feedback during development; E2E tests provide final confidence before '
    'integration.'
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — DEPENDENCY INJECTION AND MOCKING
# ══════════════════════════════════════════════════════════════════════════════
h1('Section 2 – Dependency Injection and Mocking Approach')

h2('2.1  How Constructor Injection Enables Testability')
body(
    'All service classes in the application use constructor injection exclusively. This is the pattern '
    'recommended by the Spring team and is the key design decision that makes the entire unit test suite '
    'possible without starting a Spring context.'
)

code_block(
    "@Service\n"
    "@Transactional\n"
    "public class TaskService {\n\n"
    "    private final TaskRepository taskRepository;\n"
    "    private final ProjectRepository projectRepository;\n"
    "    private final TaskMapper taskMapper;\n\n"
    "    public TaskService(TaskRepository taskRepository,\n"
    "                       ProjectRepository projectRepository,\n"
    "                       TaskMapper taskMapper) {\n"
    "        this.taskRepository = taskRepository;\n"
    "        this.projectRepository = projectRepository;\n"
    "        this.taskMapper = taskMapper;\n"
    "    }\n"
    "    // ... methods\n"
    "}"
)

body(
    'Because all dependencies are declared as parameters of the constructor, a unit test can instantiate '
    'TaskService directly by passing Mockito mock objects. No Spring container is needed. Mockito\'s '
    '@InjectMocks annotation exploits this same constructor to inject @Mock instances automatically.'
)
body(
    'Why not field injection? Field injection (@Autowired TaskRepository repo;) would make this '
    'impossible without either starting the full Spring context or using brittle Java reflection to '
    'force-inject mock values. Constructor injection makes dependencies explicit, immutable, and '
    'directly testable.'
)

h2('2.2  Where Mocks Were Used and Why')
add_table(
    ['Test Class', 'What is Mocked', 'Why'],
    [
        ['TaskServiceTest',
         'TaskRepository, ProjectRepository, TaskMapper',
         'Isolate TaskService business logic from persistence. Control exact return values '
         '(e.g. Optional.empty() to trigger 404 logic, thenAnswer to simulate entity save).'],
        ['ProjectServiceTest',
         'ProjectRepository, TaskRepository, ProjectMapper',
         'Test delete-blocking logic without a real database. Stub countByProjectIdAndStatus '
         'to return 0 or 2 to exercise both branches of the guard clause.'],
        ['TaskControllerWebMvcTest',
         'TaskService (@MockBean)',
         'Isolate the HTTP layer. Test controller annotations, request deserialization, '
         'bean validation, and GlobalExceptionHandler independently of service logic.'],
    ]
)

body(
    'ArgumentCaptor was used in TaskServiceTest.updateTaskStatus_toDone_setsCompletedAtTimestamp() to '
    'capture the Task entity passed to taskRepository.save(). This verifies that the service sets '
    'completedAt before persisting — a side-effect that cannot be observed from the method return value alone.'
)

h2('2.3  Where Mocks Were Avoided and Why')
add_table(
    ['Test Class', 'Mock Avoided', 'Why'],
    [
        ['TaskMapperTest',
         'None – no dependencies',
         'TaskMapper is a pure, stateless function with no collaborators. Direct instantiation '
         'with new TaskMapper() is the correct approach. Zero framework overhead.'],
        ['TaskRepositoryIntegrationTest',
         'No mocks – real H2 DB used',
         'The purpose is to verify that derived query methods produce correct SQL against a '
         'real schema. Mocking the repository would make the test meaningless.'],
        ['TaskManagerE2ETest',
         'No mocks – full stack',
         'The purpose is to verify all layers work together end-to-end. Mocking any layer '
         'would undermine the guarantee these tests are meant to provide.'],
    ]
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — MAPPING TO TEST PYRAMID
# ══════════════════════════════════════════════════════════════════════════════
h1('Section 3 – Mapping to Test Pyramid')

h2('3.1  Unit Tests (Bottom Layer) – 15 tests')
body(
    'Unit tests are annotated with @ExtendWith(MockitoExtension.class). No Spring context is loaded. '
    'All dependencies are replaced with Mockito @Mock instances injected via @InjectMocks.'
)

h3('TaskServiceTest (8 tests)')
add_table(
    ['Test Method', 'What It Verifies', 'Category'],
    [
        ['createTask_whenProjectExistsAndDueDateFuture_returnsTaskDTO',
         'Happy path: task created with OPEN status when project exists and due date is in the future',
         'Happy path'],
        ['createTask_whenProjectNotFound_throwsResourceNotFoundException',
         'Missing project triggers ResourceNotFoundException before any save is called',
         'Boundary'],
        ['createTask_whenDueDateIsInPast_throwsBusinessRuleException',
         'Business rule: due dates in the past are rejected',
         'Business rule'],
        ['updateTaskStatus_toDone_setsCompletedAtTimestamp',
         'Side-effect: completedAt is populated and status is DONE on the entity passed to save()',
         'Side-effect'],
        ['updateTaskStatus_fromDoneToOpen_throwsBusinessRuleException',
         'Business rule: completed tasks cannot be reopened to OPEN',
         'Business rule'],
        ['getTaskById_whenFound_returnsDTO',
         'Happy path: found task is correctly mapped to DTO',
         'Happy path'],
        ['getTaskById_whenNotFound_throwsResourceNotFoundException',
         'Missing task triggers ResourceNotFoundException',
         'Boundary'],
        ['deleteTask_whenTaskExists_callsRepositoryDelete',
         'Happy path: deleteById is called with the correct ID',
         'Delegation'],
    ],
    col_widths=[6.5, 7, 3]
)

body(
    'A representative test showing the Arrange-Act-Assert structure used throughout TaskServiceTest:'
)
code_block(
    "@Test\n"
    "void createTask_whenProjectExistsAndDueDateFuture_returnsTaskDTO() {\n"
    "    // Arrange\n"
    "    Project project = new Project(\"Alpha\", \"Alpha project\");\n"
    "    project.setId(1L);\n"
    "    CreateTaskRequest request = new CreateTaskRequest(\n"
    "            \"Implement login\", \"OAuth2 login page\",\n"
    "            TaskPriority.HIGH, LocalDate.now().plusDays(7), 1L);\n"
    "    Task savedTask = new Task(\"Implement login\", \"OAuth2 login page\",\n"
    "            TaskStatus.OPEN, TaskPriority.HIGH, LocalDate.now().plusDays(7), project);\n"
    "    savedTask.setId(10L);\n"
    "    TaskDTO expectedDTO = new TaskDTO(10L, \"Implement login\", \"OAuth2 login page\",\n"
    "            TaskStatus.OPEN, TaskPriority.HIGH, LocalDate.now().plusDays(7),\n"
    "            null, 1L, \"Alpha\");\n\n"
    "    when(projectRepository.findById(1L)).thenReturn(Optional.of(project));\n"
    "    when(taskRepository.save(any(Task.class))).thenReturn(savedTask);\n"
    "    when(taskMapper.toDTO(savedTask)).thenReturn(expectedDTO);\n\n"
    "    // Act\n"
    "    TaskDTO result = taskService.createTask(request);\n\n"
    "    // Assert\n"
    "    assertThat(result.getTitle()).isEqualTo(\"Implement login\");\n"
    "    assertThat(result.getStatus()).isEqualTo(TaskStatus.OPEN);\n"
    "    verify(taskRepository).save(any(Task.class));\n"
    "}"
)

h3('ProjectServiceTest (5 tests)')
add_table(
    ['Test Method', 'What It Verifies'],
    [
        ['createProject_whenNameIsUnique_returnsProjectDTO',
         'Happy path: unique name allows project creation'],
        ['createProject_whenNameIsDuplicate_throwsBusinessRuleException',
         'Business rule: duplicate project names are rejected'],
        ['deleteProject_whenHasInProgressTasks_throwsBusinessRuleException',
         'Business rule: project with in-progress tasks cannot be deleted'],
        ['deleteProject_whenNoInProgressTasks_callsRepositoryDelete',
         'Happy path: project with no in-progress tasks is deleted successfully'],
        ['getProjectById_whenNotFound_throwsResourceNotFoundException',
         'Boundary: missing project throws ResourceNotFoundException'],
    ],
    col_widths=[8, 9]
)

h3('TaskMapperTest (2 tests)  –  plain Java, no annotations')
add_table(
    ['Test Method', 'What It Verifies'],
    [
        ['toDTO_mapsAllFieldsCorrectly',
         'All nine fields of the Task entity appear correctly in the resulting TaskDTO'],
        ['toDTO_whenCompletedAtIsNull_dtoCompletedAtIsNull',
         'Nullable completedAt field passes through as null without a NullPointerException'],
    ],
    col_widths=[8, 9]
)

page_break()

h2('3.2  Integration Tests (Middle Layer) – 9 tests')

h3('TaskRepositoryIntegrationTest (4 tests)  –  @DataJpaTest')
body(
    '@DataJpaTest loads only the JPA slice: entity manager, repositories, and an H2 in-memory schema. '
    'No web layer or service layer is present. TestEntityManager is used to persist test data directly, '
    'and entityManager.clear() is called after each persist to flush the first-level cache, ensuring '
    'subsequent queries hit the actual database rather than returning cached objects.'
)
add_table(
    ['Test Method', 'What It Verifies'],
    [
        ['findByProjectId_returnsOnlyTasksBelongingToThatProject',
         'Derived query JOIN produces the correct result set for a given project ID'],
        ['findByStatus_returnsOnlyTasksWithMatchingStatus',
         'Enum column filtering on the status column works correctly in generated SQL'],
        ['countByProjectIdAndStatus_returnsCorrectCount',
         'COUNT query with two predicates (project ID AND status) returns an accurate number'],
        ['findByPriorityAndStatusNot_returnsCriticalNonDoneTasks',
         'NOT-EQUAL operator in the derived query correctly excludes DONE tasks'],
    ],
    col_widths=[8, 9]
)

h3('TaskControllerWebMvcTest (5 tests)  –  @WebMvcTest(TaskController.class)')
body(
    '@WebMvcTest loads only the web layer: the controller, GlobalExceptionHandler, and Jackson. '
    'TaskService is replaced with a @MockBean. MockMvc performs HTTP requests without starting a real '
    'server, testing the full request → controller → response path at the web layer only.'
)
add_table(
    ['Test Method', 'What It Verifies'],
    [
        ['createTask_withValidRequest_returns201WithTaskDTO',
         'Controller maps POST body to CreateTaskRequest, calls service, returns 201 with JSON body'],
        ['createTask_withBlankTitle_returns400',
         '@NotBlank on the title field is enforced by bean validation before reaching the service'],
        ['createTask_withNullProjectId_returns400',
         '@NotNull on the projectId field is enforced by bean validation'],
        ['getTaskById_whenTaskNotFound_returns404',
         'ResourceNotFoundException thrown by the service is mapped to HTTP 404 by GlobalExceptionHandler'],
        ['updateTaskStatus_withValidStatus_returns200',
         'PATCH /tasks/{id}/status is routed correctly and returns 200 with the updated DTO'],
    ],
    col_widths=[8, 9]
)

page_break()

h2('3.3  End-to-End Tests (Top Layer) – 4 tests')

h3('TaskManagerE2ETest (4 tests)  –  @SpringBootTest(webEnvironment = RANDOM_PORT)')
body(
    'The full Spring Boot application is started on a random port. TestRestTemplate is used to make '
    'real HTTP calls over the loopback interface. An H2 in-memory database provides a clean, isolated '
    'persistence layer for each test run. These tests verify that all layers — HTTP, controller, '
    'service, repository, and database — work together correctly.'
)
add_table(
    ['Test Method', 'What It Verifies'],
    [
        ['fullWorkflow_createProjectAndTask_thenCompleteTask',
         'Full lifecycle via real HTTP: POST project → POST task → PATCH status to DONE → '
         'GET task and verify status is DONE and completedAt is not null'],
        ['createTask_forNonExistentProject_returns404',
         'Error propagation through the full stack: 404 response body when project does not exist'],
        ['deleteProject_withInProgressTask_returns409Conflict',
         'Business rule enforced end-to-end: 409 Conflict when attempting to delete a project '
         'that has tasks with IN_PROGRESS status'],
        ['getCriticalOpenTasks_returnsOnlyCriticalNonDoneTasks',
         'Filter endpoint returns only tasks with CRITICAL priority and non-DONE status from a real DB'],
    ],
    col_widths=[8, 9]
)

body(
    'E2E tests are the most expensive in the suite (~11 s combined) because each one starts the full '
    'Spring context. They are deliberately few in number, covering only the most critical workflows and '
    'error paths that cannot be verified at a lower level of the pyramid.'
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — AI INTERACTION LOG
# ══════════════════════════════════════════════════════════════════════════════
h1('Section 4 – AI Interaction Log')
body(
    'This section documents five representative interactions with an AI assistant (Claude) during the '
    'development of the test suite. Each example is categorised as: Accepted As-Is, Modified, Rejected, '
    'Gap Identified, or Constrained. The examples are chosen to illustrate both the strengths and the '
    'failure modes observed.'
)

# ── Example 1 ──────────────────────────────────────────────────────────────────
h2('4.1  Example 1: AI Output Accepted As-Is')

h3('Prompt given to AI')
quote_block(
    '"Write a @DataJpaTest integration test for TaskRepository. Test the findByPriorityAndStatusNot '
    'query method. Persist test data using TestEntityManager, clear the entity manager cache, then '
    'assert that only tasks with CRITICAL priority and non-DONE status are returned. Use AssertJ."'
)

h3('AI-generated output  (accepted without modification)')
code_block(
    "@Test\n"
    "void findByPriorityAndStatusNot_returnsCriticalNonDoneTasks() {\n"
    "    Project project = entityManager.persistAndFlush(\n"
    "            new Project(\"Critical Project\", null));\n\n"
    "    entityManager.persistAndFlush(new Task(\"Critical open\", null,\n"
    "            TaskStatus.OPEN, TaskPriority.CRITICAL, null, project));\n"
    "    entityManager.persistAndFlush(new Task(\"Critical done\", null,\n"
    "            TaskStatus.DONE, TaskPriority.CRITICAL, null, project));\n"
    "    entityManager.persistAndFlush(new Task(\"Medium open\", null,\n"
    "            TaskStatus.OPEN, TaskPriority.MEDIUM, null, project));\n\n"
    "    entityManager.clear();  // force real SQL SELECT on next query\n\n"
    "    List<Task> result = taskRepository.findByPriorityAndStatusNot(\n"
    "            TaskPriority.CRITICAL, TaskStatus.DONE);\n\n"
    "    assertThat(result).hasSize(1);\n"
    "    assertThat(result.get(0).getTitle()).isEqualTo(\"Critical open\");\n"
    "}"
)

h3('Justification for acceptance')
body(
    'The AI correctly used persistAndFlush() followed by entityManager.clear() — the proper @DataJpaTest '
    'pattern. Without clear(), Hibernate returns objects from its first-level cache rather than executing '
    'the SQL query, which would make the test pass even if the query method were completely broken. '
    'The three test entities cover all three relevant combinations (critical-open, critical-done, medium-open), '
    'providing a thorough boundary check. The Arrange-Act-Assert structure is clean and the AssertJ '
    'assertion is typed and readable. No modification was required.'
)

# ── Example 2 ──────────────────────────────────────────────────────────────────
h2('4.2  Example 2: AI Output Modified')

h3('Original prompt')
quote_block(
    '"Write a unit test for TaskService.updateTaskStatus that verifies the completedAt timestamp '
    'is set when a task is transitioned to DONE."'
)

h3('Original AI-generated output (before modification)')
code_block(
    "@Test\n"
    "void updateTaskStatus_toDone_setsCompletedAt() {\n"
    "    Task task = new Task();\n"
    "    task.setId(5L);\n"
    "    task.setStatus(TaskStatus.IN_PROGRESS);\n\n"
    "    when(taskRepository.findById(5L)).thenReturn(Optional.of(task));\n"
    "    when(taskRepository.save(task)).thenReturn(task);\n"
    "    when(taskMapper.toDTO(task)).thenReturn(new TaskDTO());\n\n"
    "    taskService.updateTaskStatus(5L, TaskStatus.DONE);\n\n"
    "    assertThat(task.getCompletedAt()).isNotNull();  // <-- problem: asserts on local variable\n"
    "}"
)

h3('Problems identified with the original')
numbered(
    'assertThat(task.getCompletedAt()) asserts on the local reference variable, not on the entity '
    'actually passed to save(). If the service were refactored to copy the entity before saving, '
    'this assertion would pass on the original object even if the saved copy had a null completedAt.'
)
numbered(
    'new Task() with no-arg constructor leaves the project field null. When TaskMapper.toDTO() '
    'calls task.getProject().getName(), this throws a NullPointerException, causing the test '
    'to fail for the wrong reason.'
)

h3('Final modified version')
code_block(
    "@Test\n"
    "void updateTaskStatus_toDone_setsCompletedAtTimestamp() {\n"
    "    Project project = new Project(\"Gamma\", \"Gamma project\");\n"
    "    project.setId(3L);\n"
    "    Task task = new Task(\"Fix bug\", null, TaskStatus.IN_PROGRESS,\n"
    "            TaskPriority.HIGH, null, project);\n"
    "    task.setId(5L);\n\n"
    "    when(taskRepository.findById(5L)).thenReturn(Optional.of(task));\n"
    "    when(taskRepository.save(any(Task.class))).thenAnswer(inv -> inv.getArgument(0));\n"
    "    when(taskMapper.toDTO(any(Task.class))).thenReturn(\n"
    "            new TaskDTO(5L, \"Fix bug\", null, TaskStatus.DONE, TaskPriority.HIGH,\n"
    "                        null, LocalDateTime.now(), 3L, \"Gamma\"));\n\n"
    "    taskService.updateTaskStatus(5L, TaskStatus.DONE);\n\n"
    "    ArgumentCaptor<Task> captor = ArgumentCaptor.forClass(Task.class);\n"
    "    verify(taskRepository).save(captor.capture());\n"
    "    assertThat(captor.getValue().getCompletedAt()).isNotNull();\n"
    "    assertThat(captor.getValue().getStatus()).isEqualTo(TaskStatus.DONE);\n"
    "}"
)

h3('What changed and why')
bullet(
    'ArgumentCaptor introduced: captures the exact entity passed to save(), verifying the correct '
    'contract — that the service sets completedAt before persisting, not after.'
)
bullet(
    'any(Task.class) replaces the specific reference in the stub: the test no longer breaks if '
    'the service creates a new entity instance rather than mutating the original.'
)
bullet(
    'A properly constructed Project object added: prevents NullPointerException in the mapper '
    'and makes the test self-contained.'
)

page_break()

# ── Example 3 ──────────────────────────────────────────────────────────────────
h2('4.3  Example 3: AI Output Rejected')

h3('Prompt')
quote_block(
    '"Write a unit test for ProjectService.deleteProject that verifies the method throws '
    'BusinessRuleException when the project has in-progress tasks."'
)

h3('AI-generated output (rejected)')
code_block(
    "@SpringBootTest\n"
    "@Transactional\n"
    "class ProjectServiceDeleteTest {\n\n"
    "    @Autowired\n"
    "    private ProjectService projectService;\n\n"
    "    @Autowired\n"
    "    private ProjectRepository projectRepository;\n\n"
    "    @Autowired\n"
    "    private TaskRepository taskRepository;\n\n"
    "    @Test\n"
    "    void deleteProject_withInProgressTasks_throwsException() {\n"
    "        Project project = new Project(\"Doomed Project\", \"desc\");\n"
    "        projectRepository.save(project);\n"
    "        // ... taskRepository.save(...)\n"
    "        assertThrows(BusinessRuleException.class,\n"
    "                () -> projectService.deleteProject(project.getId()));\n"
    "    }\n"
    "}"
)

h3('Technical justification for rejection')
numbered(
    'Wrong layer of the Test Pyramid. @SpringBootTest boots the entire application — JPA schema '
    'creation, all beans, HTTP infrastructure. The business rule exists entirely within '
    'ProjectService.deleteProject(), which needs only two mock objects to test. This test takes '
    'over 5 seconds for something that should take under 1 millisecond.'
)
numbered(
    'Misuse of Spring context for a unit concern. A full application context provides no additional '
    'confidence for this test. The rule is "if countByProjectIdAndStatus > 0, throw '
    'BusinessRuleException". That logic has no dependency on a real database schema.'
)
numbered(
    'Fragile implicit contract. The test is coupled to: the test application profile, H2 schema '
    'initialisation, @Transactional rollback behaviour, and the autowiring of multiple repositories. '
    'Any misconfiguration causes the test to fail for reasons unrelated to the business rule.'
)

h3('Correct replacement used in the suite')
code_block(
    "@Test\n"
    "void deleteProject_whenHasInProgressTasks_throwsBusinessRuleException() {\n"
    "    when(projectRepository.existsById(10L)).thenReturn(true);\n"
    "    when(taskRepository.countByProjectIdAndStatus(\n"
    "            10L, TaskStatus.IN_PROGRESS)).thenReturn(2L);\n\n"
    "    assertThatThrownBy(() -> projectService.deleteProject(10L))\n"
    "            .isInstanceOf(BusinessRuleException.class)\n"
    "            .hasMessageContaining(\"in-progress\");\n\n"
    "    verify(projectRepository, never()).deleteById(any());\n"
    "}"
)
body('This unit test achieves the same verification in under 1 ms with zero infrastructure.')

# ── Example 4 ──────────────────────────────────────────────────────────────────
h2('4.4  Example 4: AI-Identified Gap or Improvement')

h3('Analysis prompt')
quote_block(
    '"Review the unit test suite for TaskService. Identify any missing test scenarios or '
    'edge cases not currently covered."'
)

h3("AI's suggestion")
quote_block(
    '"The current suite does not test deleteTask when the task does not exist. '
    'TaskService.deleteTask calls taskRepository.existsById(id) and throws '
    'ResourceNotFoundException if it returns false. The only existing deleteTask test covers '
    'the happy path (task exists, deleteById is called). A test for the not-found case would '
    'verify the guard clause and confirm that deleteById is never called."'
)

h3('Action taken – test added to the suite')
code_block(
    "@Test\n"
    "void deleteTask_whenTaskDoesNotExist_throwsResourceNotFoundException() {\n"
    "    when(taskRepository.existsById(404L)).thenReturn(false);\n\n"
    "    assertThatThrownBy(() -> taskService.deleteTask(404L))\n"
    "            .isInstanceOf(ResourceNotFoundException.class)\n"
    "            .hasMessageContaining(\"404\");\n\n"
    "    verify(taskRepository, never()).deleteById(any());\n"
    "}"
)

h3('Analysis of the AI suggestion')
body(
    'The suggestion was entirely correct and valuable. The AI applied a consistent principle: every '
    'method with a guard clause (if !exists → throw) requires at least two tests — one where the '
    'guard passes and one where it throws. This is easy to miss when focusing on the happy path. '
    'The suggestion was accepted without modification and brought TaskServiceTest to its final count '
    'of 8 tests, providing complete coverage of all public methods in TaskService.'
)

page_break()

# ── Example 5 ──────────────────────────────────────────────────────────────────
h2('4.5  Example 5: AI Output Constrained by Explicit Rules')

h3('Prompt with explicit constraints')
quote_block(
    '"Write a unit test for TaskMapper.toDTO. Apply the following constraints:\n'
    '1. Do NOT use @ExtendWith(MockitoExtension.class) or any Mockito annotations.\n'
    '2. Do NOT start a Spring context (@SpringBootTest, @DataJpaTest, @WebMvcTest).\n'
    '3. Instantiate TaskMapper with new TaskMapper() directly.\n'
    '4. Use only JUnit 5 and AssertJ – no Spring test utilities.\n'
    'Verify that all fields of the Task entity are correctly mapped to TaskDTO."'
)

h3('AI-generated output produced under this constraint')
code_block(
    "class TaskMapperTest {\n\n"
    "    private final TaskMapper mapper = new TaskMapper();\n\n"
    "    @Test\n"
    "    void toDTO_mapsAllFieldsCorrectly() {\n"
    "        Project project = new Project(\"Mapper Project\", \"desc\");\n"
    "        project.setId(5L);\n\n"
    "        Task task = new Task(\"Map me\", \"description here\",\n"
    "                TaskStatus.IN_PROGRESS, TaskPriority.HIGH,\n"
    "                LocalDate.of(2026, 6, 1), project);\n"
    "        task.setId(99L);\n"
    "        task.setCompletedAt(LocalDateTime.of(2026, 5, 20, 10, 0));\n\n"
    "        TaskDTO dto = mapper.toDTO(task);\n\n"
    "        assertThat(dto.getId()).isEqualTo(99L);\n"
    "        assertThat(dto.getTitle()).isEqualTo(\"Map me\");\n"
    "        assertThat(dto.getDescription()).isEqualTo(\"description here\");\n"
    "        assertThat(dto.getStatus()).isEqualTo(TaskStatus.IN_PROGRESS);\n"
    "        assertThat(dto.getPriority()).isEqualTo(TaskPriority.HIGH);\n"
    "        assertThat(dto.getDueDate()).isEqualTo(LocalDate.of(2026, 6, 1));\n"
    "        assertThat(dto.getCompletedAt())\n"
    "                .isEqualTo(LocalDateTime.of(2026, 5, 20, 10, 0));\n"
    "        assertThat(dto.getProjectId()).isEqualTo(5L);\n"
    "        assertThat(dto.getProjectName()).isEqualTo(\"Mapper Project\");\n"
    "    }\n"
    "}"
)

h3('Why the constraint was necessary')
body(
    'Without the constraint, AI tools consistently add @ExtendWith(MockitoExtension.class) or '
    '@SpringBootTest to every test class by default — even when the class under test has zero '
    'dependencies. TaskMapper is a stateless, dependency-free class that performs only field '
    'mapping. Testing it requires nothing more than new TaskMapper() and a method call.'
)

h3('How the constraint improved test quality')
body(
    'The resulting test runs in under 1 millisecond. It has zero framework overhead and cannot be '
    'broken by database configuration, Spring Boot version changes, Mockito bytecode manipulation, '
    'or test profile settings. It tests exactly one thing: the public contract of the toDTO() method. '
    'This is a textbook example of a test correctly placed at the very base of the Test Pyramid. '
    'The constraint eliminated the most common AI failure mode before any code was generated.'
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — EVALUATION AND REFLECTION
# ══════════════════════════════════════════════════════════════════════════════
h1('Section 5 – Evaluation and Reflection')

h2('Strengths of AI-Generated Tests')

h3('Strength 1: Boilerplate generation speed')
body(
    'AI generated the structural boilerplate of test classes — imports, class-level annotations, '
    'mock field declarations, and @BeforeEach setup methods — in seconds. For '
    'TaskRepositoryIntegrationTest, the AI correctly identified @DataJpaTest, '
    '@ActiveProfiles("test"), TestEntityManager injection, and the persistAndFlush/clear pattern '
    'without being explicitly prompted. Writing this setup manually is error-prone and time-consuming; '
    'the AI produced it correctly on the first attempt, freeing the developer to focus on the actual '
    'test logic rather than infrastructure scaffolding.'
)

h3('Strength 2: Comprehensive boundary condition identification')
body(
    'When asked to review the test suite (Example 4), the AI correctly identified the missing '
    'deleteTask not-found test. It applied the principle — every guard clause needs both a passing '
    'and a failing test — consistently and without prompting. This is the kind of thorough, '
    'systematic enumeration that humans focusing on happy-path scenarios can miss, particularly '
    'when working under time pressure. The AI served effectively as a second reviewer.'
)

h3('Strength 3: Knowledge of framework-specific patterns')
body(
    'The AI demonstrated accurate knowledge of Spring test slices. It correctly distinguished between '
    'scenarios requiring @DataJpaTest (JPA-only slice), @WebMvcTest (web-only slice), and '
    '@SpringBootTest (full context), and applied the right annotation without being told which to use '
    'in Examples 1 and the integration test cases. This knowledge of how to minimise Spring context '
    'loading is non-trivial for beginners and was produced correctly by the AI.'
)

h2('Weaknesses of AI-Generated Tests')

h3('Weakness 1: Default to heavy test infrastructure')
body(
    'Example 3 demonstrates the most consistent AI failure mode: reaching for @SpringBootTest '
    'regardless of whether the scenario requires it. The prompt explicitly asked for a "unit test" '
    'for a service method, yet the AI generated a full @SpringBootTest class. Left unconstrained, '
    'AI tools pattern-match "Spring Boot application" → "use @SpringBootTest" without evaluating '
    'whether the specific scenario has any dependency on the full context. The practical consequence '
    'is test suites that are slow, brittle, and fragile under configuration changes — for reasons '
    'entirely unrelated to the business logic under test.'
)

h3('Weakness 2: Testing implementation details rather than contracts')
body(
    'In Example 2, the original AI-generated test asserted on the local task variable rather than '
    'using ArgumentCaptor to inspect the entity passed to save(). The test appeared correct and would '
    'have passed in all foreseeable circumstances — but it verified Java reference semantics rather '
    'than the actual contract: that the service sets completedAt before persisting. This is a subtle '
    'but important distinction. A test that passes for the wrong reason can survive a refactoring that '
    'breaks the actual contract, providing false confidence.'
)

h2('Observed AI Failure Mode: Incorrect Assumption About Test Scope')
body(
    'Across multiple interactions, AI tools consistently pattern-matched on "Spring Boot application" '
    'and defaulted to @SpringBootTest for every test scenario, regardless of whether the scenario '
    'required the full context. This is not an occasional error — it is the default behaviour. The '
    'implication is that without human oversight, an AI-generated test suite for a Spring application '
    'would consist almost entirely of slow, infrastructure-coupled @SpringBootTest tests. The correct '
    'pyramid distribution (54% unit / 32% integration / 14% E2E) requires active constraint and '
    'correction of AI output.'
)

h2('Example Where Human Judgement Improved Test Quality')
body(
    'In Example 2, the AI\'s assertion on the mutated local variable appeared correct at first glance — '
    'the completedAt field was being set, and the assertion was checking it. It was human review of the '
    'reason for the assertion, not just the assertion itself, that identified the problem. The key '
    'question was: "Is this test verifying the contract (that the entity passed to save() has '
    'completedAt set), or is it just verifying Java reference behaviour (that mutating the original '
    'object also mutates the variable pointing to it)?" AI does not naturally ask this kind of '
    'reflective question about its own output. The fix — introducing ArgumentCaptor — came from '
    'applying the testing principle that assertions should verify observable outputs through public '
    'interfaces, not through internal object state.'
)

h2('One Improvement to the AI-Assisted Testing Approach')
body(
    'The most valuable single improvement would be to establish a constraint prompt template '
    'prepended to every test-generation request. Based on patterns observed in this assignment, '
    'the template would include:'
)
code_block(
    "Rules for all tests in this project:\n"
    "1. If the class under test has no Spring dependencies, use\n"
    "   @ExtendWith(MockitoExtension.class) only. Do not start a Spring context.\n"
    "2. Never use @SpringBootTest for a test that can be written with @WebMvcTest,\n"
    "   @DataJpaTest, or plain Mockito.\n"
    "3. When asserting on side-effects (calls to save() or delete()), use\n"
    "   ArgumentCaptor to verify the exact argument passed to the repository method.\n"
    "4. Every method with a guard clause requires at least two tests:\n"
    "   one where the guard passes and one where it throws."
)
body(
    'Rather than correcting AI output after the fact, this approach constrains the AI at the source '
    '— eliminating the most frequently observed failure modes before any code is generated. The '
    'template embodies the lessons learned from Examples 3, 2, and 4 respectively.'
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
h1('Conclusion')
body(
    'This assignment set out to build a principled, AI-assisted test suite for a Task Manager REST API '
    'and to reflect critically on that experience. The final suite of 28 tests — 15 unit, 9 integration, '
    '4 end-to-end — achieves a distribution close to the recommended Test Pyramid proportions. All '
    'business rules are covered at the unit level, all HTTP contracts are verified at the integration '
    'level, and the most critical end-to-end workflows are confirmed at the E2E level.'
)
body(
    'The AI assistant proved genuinely useful for two specific tasks: generating boilerplate '
    'infrastructure for test classes (especially @DataJpaTest setup) and systematically identifying '
    'missing boundary-condition tests. In both cases, the AI produced correct output that would have '
    'taken significantly more time to write manually.'
)
body(
    'However, the AI also demonstrated consistent failure modes that required human correction. The '
    'tendency to default to @SpringBootTest for every test scenario is the most significant — if left '
    'uncorrected, it would have produced a suite dominated by slow, infrastructure-coupled tests sitting '
    'at the apex of the pyramid rather than the base. The second failure mode — asserting on implementation '
    'details (local variable state) rather than observable contracts (the entity passed to save()) — '
    'is subtler but equally important.'
)
body(
    'The key takeaway is that AI-assisted testing is most effective as a collaboration: the AI handles '
    'speed and breadth (boilerplate, enumeration of test cases), while the human provides judgement '
    'about test placement, assertion correctness, and the question of what the test is actually proving. '
    'Neither alone produces a high-quality test suite; both together can.'
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════════════════════════
h1('References')

refs = [
    ('Cohn, M. (2009)',
     'Succeeding with Agile: Software Development Using Scrum. Addison-Wesley Professional.'),
    ('Fowler, M. (2012)',
     'TestPyramid. martinfowler.com. Available at: https://martinfowler.com/bliki/TestPyramid.html'),
    ('Spring Framework Documentation (2024)',
     'Testing. Spring Boot Reference Documentation. '
     'Available at: https://docs.spring.io/spring-boot/docs/current/reference/html/features.html#features.testing'),
    ('JUnit 5 Team (2024)',
     'JUnit 5 User Guide. Available at: https://junit.org/junit5/docs/current/user-guide/'),
    ('Mockito Project (2024)',
     'Mockito Documentation. Available at: https://javadoc.io/doc/org.mockito/mockito-core/latest/'),
    ('AssertJ Team (2024)',
     'AssertJ – Fluent Assertions for Java. Available at: https://assertj.github.io/doc/'),
    ('Freeman, S. & Pryce, N. (2009)',
     'Growing Object-Oriented Software, Guided by Tests. Addison-Wesley Professional.'),
]

for i, (author, detail) in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.8)
    rb = p.add_run(f'[{i}]  {author}. ')
    rb.bold = True
    rb.font.size = Pt(10.5)
    rv = p.add_run(detail)
    rv.font.size = Pt(10.5)

spacer(20)

# ── Footer note ────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
r = p.add_run('GitHub Repository: https://github.com/CodingKirthi/CBD-Taskmanager')
r.font.size = Pt(9.5)
r.font.color.rgb = MID_BLUE
r.italic = True

# ── Save ───────────────────────────────────────────────────────────────────────
doc.save('c:/CBD/CBD-Assignment-2/SUBMISSION_REPORT.docx')
print('SUBMISSION_REPORT.docx saved successfully.')
