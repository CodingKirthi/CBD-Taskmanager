"""
Generates DEMO_SCRIPT.docx - Screencast demo guide for CBD Assignment 2.
Run: python make_demo_script.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Colours ───────────────────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x1F, 0x38, 0x64)
MID_BLUE   = RGBColor(0x2F, 0x54, 0x96)
GREEN      = RGBColor(0x37, 0x86, 0x44)
ORANGE     = RGBColor(0xC5, 0x5A, 0x11)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

# ── Helpers ───────────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_paragraph()
    p.clear()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = DARK_BLUE
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(6)
    return p

def h2(text):
    p = doc.add_paragraph()
    p.clear()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = MID_BLUE
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p

def h3(text, color=None):
    p = doc.add_paragraph()
    p.clear()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = color or MID_BLUE
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    return p

def body(text):
    p = doc.add_paragraph()
    p.add_run(text)
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(2)
    return p

def say_block(text):
    """Blue italic box – what to SAY during demo."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run('"' + text + '"')
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = MID_BLUE
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EBF3FB')
    pPr.append(shd)
    return p

def code_block(text):
    """Grey box for commands / JSON."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    return p

def label(text, color):
    """Small bold coloured label e.g. WHAT TO SHOW / WHAT TO SAY."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(2)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = WHITE
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F3864')
        tcPr.append(shd)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('CBD Assignment 2 — Screencast Demo Script')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = DARK_BLUE

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Task Manager REST API (Spring Boot 3.2)')
run2.font.size = Pt(14)
run2.font.color.rgb = MID_BLUE

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.add_run('Follow this script step by step during your recording')

doc.add_paragraph()

# ── QUICK SUMMARY CHECKLIST ───────────────────────────────────────────────────
h1('Quick Summary Checklist')
body('Complete every item in order. Tick each one as you go.')

add_table(
    ['#', 'Step', 'Time', 'Done?'],
    [
        ['1', 'Setup — open IDE, Postman, terminal, browser', 'Before recording', '[ ]'],
        ['2', 'Introduction — show GitHub repo', '30 sec', '[ ]'],
        ['3', 'Start the application', '1 min', '[ ]'],
        ['4', 'Quick code walkthrough in IDE', '2 min', '[ ]'],
        ['5', 'API demo in Postman (11 actions)', '5 min', '[ ]'],
        ['6', 'Show H2 database console', '1 min', '[ ]'],
        ['7', 'Run all tests in terminal', '1 min', '[ ]'],
        ['8', 'Explain the 3 test levels in IDE', '3 min', '[ ]'],
        ['9', 'Show JaCoCo coverage report', '1 min', '[ ]'],
        ['10', 'Closing — back to GitHub repo', '30 sec', '[ ]'],
    ]
)

# ── BEFORE YOU START ──────────────────────────────────────────────────────────
h1('Before You Start — Setup (Do This Before Recording)')

body('Open all four of the following before you press record:')
bullet('IntelliJ IDEA (or VSCode) with the project open')
bullet('Postman — import TaskManager.postman_collection.json via Import button')
bullet('A terminal / command prompt at   c:\\CBD\\CBD-Assignment-2')
bullet('A web browser (for H2 console later)')

body('Keep a sticky note visible on screen with these blanks to fill in during the demo:')
add_table(
    ['Item', 'ID (fill in during demo)'],
    [
        ['Project ID', ''],
        ['Task 1 ID (My First Task)', ''],
        ['Task 2 ID (Fix Critical Bug)', ''],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# PART 1
# ══════════════════════════════════════════════════════════════════════════════
h1('Part 1 — Introduction (30 seconds)')

label('WHAT TO SHOW ON SCREEN:', ORANGE)
body('Your GitHub repository page: https://github.com/CodingKirthi/AI-Assisted-Testing-of-a-Spring-Boot-application')

label('WHAT TO SAY:', MID_BLUE)
say_block(
    'This is my CBD Assignment 2 — a Task Manager REST API built with Java and Spring Boot. '
    'It lets you create Projects and Tasks, manage task statuses, and enforces business rules. '
    'The application is fully tested with unit tests, integration tests, and end-to-end tests.'
)

# ══════════════════════════════════════════════════════════════════════════════
# PART 2
# ══════════════════════════════════════════════════════════════════════════════
h1('Part 2 — Start the Application (1 minute)')

label('WHAT TO DO:', GREEN)
body('In your terminal, run:')
code_block('./mvnw spring-boot:run')
body('Wait until you see this line in the console output:')
code_block('Started TaskManagerApplication in X.XXX seconds')

label('WHAT TO SAY:', MID_BLUE)
say_block(
    'I am starting the Spring Boot application. It connects to an H2 file-based database '
    'and listens on port 8080. Once started, all API endpoints are live and ready.'
)

# ══════════════════════════════════════════════════════════════════════════════
# PART 3
# ══════════════════════════════════════════════════════════════════════════════
h1('Part 3 — Quick Code Walkthrough (2 minutes)')

h2('Step 3.1 — Open Task.java')
label('FILE:', ORANGE)
body('src/main/java/com/university/taskmanager/model/Task.java')
label('WHAT TO SAY:', MID_BLUE)
say_block(
    'This is the Task model. Each task has a title, description, a status which can be OPEN, '
    'IN_PROGRESS, or DONE, a priority from LOW to CRITICAL, an optional due date, and it '
    'belongs to a Project.'
)

h2('Step 3.2 — Open TaskService.java')
label('FILE:', ORANGE)
body('src/main/java/com/university/taskmanager/service/TaskService.java')
label('POINT TO:', GREEN)
bullet('Line 42 — due date cannot be in the past')
bullet('Line 80 — DONE task cannot be reopened to OPEN')
bullet('Line 86 — completedAt is automatically recorded when status becomes DONE')
label('WHAT TO SAY:', MID_BLUE)
say_block(
    'The Service layer enforces business rules. Line 42 rejects any task with a due date in the past. '
    'Line 80 prevents reopening a completed task. And when a task is marked DONE, '
    'the exact timestamp is automatically recorded in completedAt.'
)

h2('Step 3.3 — Open TaskController.java')
label('FILE:', ORANGE)
body('src/main/java/com/university/taskmanager/controller/TaskController.java')
label('WHAT TO SAY:', MID_BLUE)
say_block(
    'The Controller exposes the API endpoints. Each method maps to an HTTP method and URL. '
    'POST /api/tasks creates a task, and PATCH /api/tasks/{id}/status updates its status.'
)

# ══════════════════════════════════════════════════════════════════════════════
# PART 4
# ══════════════════════════════════════════════════════════════════════════════
h1('Part 4 — API Demo in Postman (5 minutes)')

body('Open the "Task Manager API" collection in Postman. Run each action in the order below.')

# --- Action 1 ---
h2('Action 1 — Create a Project')
label('POSTMAN REQUEST:', ORANGE)
body('Projects  →  Create Project  →  Send')
code_block('{\n  "name": "My First Project",\n  "description": "A project to manage my tasks"\n}')
label('EXPECTED RESPONSE:', GREEN)
body('201 Created — note the id field in the response and write it on your sticky note.')
label('WHAT TO SAY:', MID_BLUE)
say_block(
    'I am creating a project. The API returns 201 Created and gives back the project '
    'with an auto-generated ID. I will note this ID and use it when creating tasks.'
)

# --- Action 2 ---
h2('Action 2 — Try to Create a Duplicate Project')
label('POSTMAN REQUEST:', ORANGE)
body('Projects  →  Create Duplicate Project (409 Error Demo)  →  Send')
label('EXPECTED RESPONSE:', GREEN)
body('409 Conflict  —  message: "A project with name My First Project already exists"')
label('WHAT TO SAY:', MID_BLUE)
say_block(
    'If I try to create another project with the same name the system blocks it and returns '
    '409 Conflict. This is a business rule enforced inside ProjectService.'
)

# --- Action 3 ---
h2('Action 3 — Create a Task')
label('POSTMAN REQUEST:', ORANGE)
body('Tasks  →  Create Task  →  check projectId matches your sticky note  →  Send')
code_block('{\n  "title": "My First Task",\n  "description": "A task for my project",\n  "priority": "HIGH",\n  "dueDate": "2026-12-31",\n  "projectId": 1\n}')
label('EXPECTED RESPONSE:', GREEN)
body('201 Created — note the task id and write it on your sticky note.')
label('WHAT TO SAY:', MID_BLUE)
say_block(
    'I create a task inside that project. It starts with status OPEN by default. '
    'The priority is HIGH and the due date is in the future.'
)

# --- Action 4 ---
h2('Action 4 — Past Due Date Validation')
label('POSTMAN REQUEST:', ORANGE)
body('Tasks  →  Create Task - Past Due Date (400 Error Demo)  →  Send')
label('EXPECTED RESPONSE:', GREEN)
body('409 Conflict  —  message: "Due date cannot be in the past"')
label('WHAT TO SAY:', MID_BLUE)
say_block(
    'If someone sends a due date in the past — like the year 2020 — the service rejects it '
    'before anything is saved to the database.'
)

# --- Action 5 ---
h2('Action 5 — Blank Title Validation')
label('POSTMAN REQUEST:', ORANGE)
body('Tasks  →  Create Task - Missing Title (400 Error Demo)  →  Send')
label('EXPECTED RESPONSE:', GREEN)
body('400 Bad Request')
label('WHAT TO SAY:', MID_BLUE)
say_block(
    'A blank title is caught by Spring Bean Validation before it even reaches the service, '
    'returning a 400 Bad Request immediately.'
)

# --- Action 6 ---
h2('Action 6 — Create a Critical Task')
label('POSTMAN REQUEST:', ORANGE)
body('Tasks  →  Create Critical Task  →  Send')
code_block('{\n  "title": "Fix Critical Bug",\n  "priority": "CRITICAL",\n  "projectId": 1\n}')
label('EXPECTED RESPONSE:', GREEN)
body('201 Created — note this task id on your sticky note.')
label('WHAT TO SAY:', MID_BLUE)
say_block(
    'I create a second task with CRITICAL priority. I will use this shortly to '
    'demonstrate the critical tasks filter.'
)

# --- Action 7 ---
h2('Action 7 — Move Task 1 to IN_PROGRESS')
label('POSTMAN REQUEST:', ORANGE)
body('Tasks  →  Update Status → IN_PROGRESS  →  check URL has Task 1 ID  →  Send')
code_block('PATCH /api/tasks/1/status?status=IN_PROGRESS')
label('EXPECTED RESPONSE:', GREEN)
body('200 OK  —  "status": "IN_PROGRESS"')
label('WHAT TO SAY:', MID_BLUE)
say_block(
    'I move the first task to IN_PROGRESS. The status field in the response has updated.'
)

# --- Action 8 ---
h2('Action 8 — Try to Delete the Project (Blocked!)')
label('POSTMAN REQUEST:', ORANGE)
body('Projects  →  Delete Project  →  check URL has your Project ID  →  Send')
label('EXPECTED RESPONSE:', GREEN)
body('409 Conflict  —  message: "Cannot delete project with 1 in-progress task(s)"')
label('WHAT TO SAY:', MID_BLUE)
say_block(
    'You cannot delete a project that has tasks still in progress. '
    'The system protects data integrity and returns a clear 409 Conflict error.'
)

# --- Action 9 ---
h2('Action 9 — Get Critical Open Tasks')
label('POSTMAN REQUEST:', ORANGE)
body('Tasks  →  Get Critical Open Tasks  →  Send')
label('EXPECTED RESPONSE:', GREEN)
body('200 OK — list containing only the CRITICAL task that is not DONE (Fix Critical Bug)')
label('WHAT TO SAY:', MID_BLUE)
say_block(
    'This endpoint returns only CRITICAL priority tasks that have not been completed. '
    'Notice it returns only the Fix Critical Bug task and not the HIGH priority one.'
)

# --- Action 10 ---
h2('Action 10 — Mark Task 1 as DONE')
label('POSTMAN REQUEST:', ORANGE)
body('Tasks  →  Update Status → DONE  →  check URL has Task 1 ID  →  Send')
label('EXPECTED RESPONSE:', GREEN)
body('200 OK  —  "status": "DONE"  and  "completedAt": "2026-..."  (a timestamp is now set)')
label('WHAT TO SAY:', MID_BLUE)
say_block(
    'When I mark a task as DONE, the system automatically records the exact date and time '
    'it was completed in the completedAt field. Also — once DONE, you cannot move it back '
    'to OPEN. The system blocks that too.'
)

# --- Action 11 ---
h2('Action 11 — Get All Tasks for the Project')
label('POSTMAN REQUEST:', ORANGE)
body('Tasks  →  Get Tasks by Project  →  check URL has your Project ID  →  Send')
label('EXPECTED RESPONSE:', GREEN)
body('200 OK — list showing both tasks')
label('WHAT TO SAY:', MID_BLUE)
say_block(
    'This shows all tasks belonging to our project. You can see the completed task '
    'with its timestamp and the critical task still open.'
)

# ══════════════════════════════════════════════════════════════════════════════
# PART 5
# ══════════════════════════════════════════════════════════════════════════════
h1('Part 5 — Show the H2 Database Console (1 minute)')

label('WHAT TO DO:', GREEN)
body('Open your browser and go to:')
code_block('http://localhost:8080/h2-console')
body('Fill in the login form:')
bullet('JDBC URL:   jdbc:h2:file:./taskmanager-db')
bullet('Username:   sa')
bullet('Password:   (leave blank)')
body('Click Connect, then run these SQL queries one at a time:')
code_block('SELECT * FROM TASKS;\nSELECT * FROM PROJECTS;')

label('WHAT TO SAY:', MID_BLUE)
say_block(
    'The H2 console lets us look directly into the database. '
    'You can see all projects and tasks are persisted, including the completedAt '
    'timestamp for the task we just completed.'
)

# ══════════════════════════════════════════════════════════════════════════════
# PART 6
# ══════════════════════════════════════════════════════════════════════════════
h1('Part 6 — Run the Tests (1 minute)')

label('WHAT TO DO:', GREEN)
body('In your terminal run:')
code_block('./mvnw test')
body('Wait for the build to finish. You should see:')
code_block('Tests run: 28, Failures: 0, Errors: 0\nBUILD SUCCESS')

label('WHAT TO SAY:', MID_BLUE)
say_block(
    'Now I will run the full test suite. The project has 28 tests across three levels — '
    'unit, integration, and end-to-end. All 28 tests pass with zero failures.'
)

# ══════════════════════════════════════════════════════════════════════════════
# PART 7
# ══════════════════════════════════════════════════════════════════════════════
h1('Part 7 — Explain the 3 Test Levels in IDE (3 minutes)')

h2('Step 7.1 — Unit Tests')
label('FILE TO OPEN:', ORANGE)
body('src/test/java/com/university/taskmanager/unit/TaskServiceTest.java')
label('POINT TO:', GREEN)
bullet('@Mock annotations at the top — these are fake versions of the database')
bullet('@InjectMocks — this wires the fakes into TaskService automatically')
bullet('Line 102 — test for past due date throws BusinessRuleException')
bullet('Line 154 — test that a DONE task cannot be reopened')
label('WHAT TO SAY:', MID_BLUE)
say_block(
    'Unit tests test a single class in complete isolation. Here I am testing TaskService. '
    'The @Mock annotations create fake versions of the database — no real database is touched. '
    'This makes each test run in under one millisecond. '
    'For example, this test checks that a past due date throws a BusinessRuleException.'
)

h2('Step 7.2 — Integration Tests (Web Layer)')
label('FILE TO OPEN:', ORANGE)
body('src/test/java/com/university/taskmanager/integration/TaskControllerWebMvcTest.java')
label('POINT TO:', GREEN)
bullet('@WebMvcTest at the top — loads only the HTTP layer, not the full application')
bullet('@MockBean TaskService — the service is still faked')
bullet('Line 69 — blank title test returns 400 Bad Request')
bullet('Line 98 — task not found returns 404 via GlobalExceptionHandler')
label('WHAT TO SAY:', MID_BLUE)
say_block(
    'Integration tests test how the Controller and HTTP layer work together. '
    '@WebMvcTest loads only the web layer. MockMvc simulates HTTP requests without '
    'starting a real server. For example, line 69 checks that a blank title returns '
    'a 400 Bad Request response.'
)

h2('Step 7.3 — End-to-End Tests')
label('FILE TO OPEN:', ORANGE)
body('src/test/java/com/university/taskmanager/e2e/TaskManagerE2ETest.java')
label('POINT TO:', GREEN)
bullet('@SpringBootTest — the full application starts on a real random port')
bullet('TestRestTemplate — makes real HTTP calls')
bullet('Line 60 — full workflow test matching exactly what we did in Postman')
bullet('Line 115 — the delete blocked by in-progress task test')
label('WHAT TO SAY:', MID_BLUE)
say_block(
    'End-to-end tests test everything together. The full application starts on a real port, '
    'TestRestTemplate makes real HTTP calls, and a real H2 database is used. '
    'Line 60 runs the same workflow I just demonstrated in Postman — '
    'create project, create task, mark DONE, then verify completedAt was set.'
)

# ══════════════════════════════════════════════════════════════════════════════
# PART 8
# ══════════════════════════════════════════════════════════════════════════════
h1('Part 8 — JaCoCo Coverage Report (1 minute)')

label('WHAT TO DO:', GREEN)
body('Open this file in your browser (drag and drop it):')
code_block('c:\\CBD\\CBD-Assignment-2\\target\\site\\jacoco\\index.html')

label('WHAT TO SAY:', MID_BLUE)
say_block(
    'JaCoCo is a tool that measures how much of the source code is exercised by tests. '
    'The green lines are covered. This report gives a percentage for each class. '
    'High coverage means our tests are thorough and most of the code has been verified.'
)

# ══════════════════════════════════════════════════════════════════════════════
# PART 9
# ══════════════════════════════════════════════════════════════════════════════
h1('Part 9 — Closing (30 seconds)')

label('WHAT TO SHOW:', ORANGE)
body('Switch back to the GitHub repository page.')

label('WHAT TO SAY:', MID_BLUE)
say_block(
    'To summarise — I built a Task Manager REST API with Spring Boot that enforces '
    'real business rules such as preventing past due dates and blocking deletion of '
    'active projects. The application is tested at three levels: unit, integration, '
    'and end-to-end. All 28 tests pass and the code is version controlled on GitHub.'
)

# ══════════════════════════════════════════════════════════════════════════════
# QUICK REFERENCE CARD
# ══════════════════════════════════════════════════════════════════════════════
h1('Quick Reference — All Postman Actions')

add_table(
    ['#', 'Postman Request Name', 'Expected Status', 'Key Point to Say'],
    [
        ['1',  'Create Project',                       '201 Created',  'Note the project ID'],
        ['2',  'Create Duplicate Project (409 Demo)',  '409 Conflict', 'Business rule: unique names'],
        ['3',  'Create Task',                          '201 Created',  'Note the task ID, starts OPEN'],
        ['4',  'Create Task - Past Due Date (Demo)',   '409 Conflict', 'Business rule: no past dates'],
        ['5',  'Create Task - Missing Title (Demo)',   '400 Bad Request','Validation before service'],
        ['6',  'Create Critical Task',                 '201 Created',  'Note this task ID too'],
        ['7',  'Update Status → IN_PROGRESS',          '200 OK',       'Status changed'],
        ['8',  'Delete Project',                       '409 Conflict', 'Business rule: in-progress blocked'],
        ['9',  'Get Critical Open Tasks',              '200 OK',       'Only CRITICAL + not DONE shown'],
        ['10', 'Update Status → DONE',                 '200 OK',       'completedAt is now set!'],
        ['11', 'Get Tasks by Project',                 '200 OK',       'Both tasks visible'],
    ]
)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save('c:/CBD/CBD-Assignment-2/DEMO_SCRIPT.docx')
print('DEMO_SCRIPT.docx saved successfully.')
